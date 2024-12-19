from docx import Document

def organize_and_save_docx(file_path):
    # Read the content of the .docx file
    doc = Document(file_path)
    organized_data = {}
    current_heading = "General"  # Default heading if no headings are found

    # Extract and group content
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Check if the paragraph is a heading
        if paragraph.style.name.startswith("Heading"):
            current_heading = text
            organized_data[current_heading] = []
        else:
            organized_data.setdefault(current_heading, []).append(text)

    # Create a new document for organized content
    new_doc = Document()

    # Rewrite organized data into the file
    for heading, paragraphs in organized_data.items():
        new_doc.add_heading(heading, level=2)
        for paragraph in paragraphs:
            new_doc.add_paragraph(paragraph)
        new_doc.add_paragraph()  # Add some spacing between sections

    # Save back to the original file (or a new file if preferred)
    new_doc.save(file_path)
    print(f"File '{file_path}' has been organized and saved.")

# Example usage
organize_and_save_docx("Ancient_Egypt.docx")

